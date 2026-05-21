"""DOCX 解析 — python-docx"""
import hashlib
from pathlib import Path

from app.models.unified_document import (
    UnifiedDocument, TextElement, ImageElement,
    TableElement, BoundingBox,
)


class DocxParser:

    async def parse(self, path: Path, session_id: str) -> UnifiedDocument:
        from docx import Document

        doc = Document(str(path))
        result = UnifiedDocument(
            source_file=str(path),
            source_format="docx",
            parse_method="python-docx",
        )

        assets_dir = path.parent / "assets"
        assets_dir.mkdir(exist_ok=True)

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            level = self._heading_level(para)
            result.texts.append(TextElement(
                id=f"docx_p{idx}",
                content=text,
                page=0,
                level=level,
                style=para.style.name if para.style else "Normal",
                fingerprint=hashlib.md5(text.encode()).hexdigest(),
            ))

        img_idx = 0
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue

            try:
                img_bytes = rel.target_part.blob
                ext = rel.target_part.content_type.split("/")[-1]
                if ext == "jpeg":
                    ext = "jpg"
                img_hash = hashlib.md5(img_bytes).hexdigest()[:12]

                img_path = assets_dir / f"docx_img{img_idx}_{img_hash}.{ext}"
                img_path.write_bytes(img_bytes)

                result.images.append(ImageElement(
                    id=f"docx_img{img_idx}_{img_hash}",
                    local_path=str(img_path),
                    page=0,
                    hash=img_hash,
                ))
                img_idx += 1
            except Exception:
                result.parse_warnings.append(f"图片提取失败 rel={rel.rId}")

        for tbl_idx, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(row_data)

            result.tables.append(TableElement(
                id=f"docx_tbl{tbl_idx}",
                page=0,
                data=data,
                headers=data[0] if data else [],
            ))

        return result

    @staticmethod
    def _heading_level(paragraph) -> int:
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if "heading" in style_name or "标题" in style_name:
            for i in range(1, 7):
                if str(i) in style_name:
                    return i
            return 1
        return 0
