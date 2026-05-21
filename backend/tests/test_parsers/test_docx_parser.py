"""Unit tests for DOCX parser."""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.parsers.docx_parser import DocxParser
from app.models.unified_document import UnifiedDocument


class TestDocxParser:
    """Test suite for DOCX parser."""

    @pytest.fixture
    def sample_docx_with_content(self, tmp_path: Path) -> Path:
        """Create a sample DOCX file with heading, paragraph, and table."""
        doc = Document()

        # Add heading level 1
        heading1 = doc.add_heading('Main Heading', level=1)
        heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add heading level 2
        heading2 = doc.add_heading('Sub Heading', level=2)

        # Add paragraph
        paragraph = doc.add_paragraph('This is a sample paragraph with some text content.')

        # Add heading level 3
        heading3 = doc.add_heading('Sub-sub Heading', level=3)

        # Add another paragraph
        doc.add_paragraph('Second paragraph for testing.')

        # Add table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # Populate table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Column 1'
        hdr_cells[1].text = 'Column 2'

        # Populate table data
        row_cells = table.rows[1].cells
        row_cells[0].text = 'Data 1'
        row_cells[1].text = 'Data 2'

        row_cells = table.rows[2].cells
        row_cells[0].text = 'Data 3'
        row_cells[1].text = 'Data 4'

        # Save file
        file_path = tmp_path / "sample.docx"
        doc.save(file_path)
        return file_path

    @pytest.fixture
    def empty_docx(self, tmp_path: Path) -> Path:
        """Create an empty DOCX file."""
        doc = Document()

        file_path = tmp_path / "empty.docx"
        doc.save(file_path)
        return file_path

    @pytest.mark.asyncio
    async def test_parser_docx_with_heading_and_table_returns_unified_document(
        self,
        sample_docx_with_content: Path
    ) -> None:
        """Test parsing DOCX with heading, paragraph and table returns UnifiedDocument."""
        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(sample_docx_with_content)

        # Verify type
        assert isinstance(result, UnifiedDocument)

        # Verify texts extracted
        assert len(result.texts) > 0, "Should extract text from DOCX"

        # Verify tables extracted
        assert len(result.tables) > 0, "Should extract table"

        # Verify fingerprint generated
        assert result.fingerprint != "", "Should generate fingerprint"

        # Verify parse method
        assert result.parse_method == "python-docx", "Should use python-docx parser"

    @pytest.mark.asyncio
    async def test_parser_docx_heading_level_1_recognition(self, sample_docx_with_content: Path) -> None:
        """Test that heading level 1 is correctly recognized."""
        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(sample_docx_with_content)

        # Find heading with level 1
        heading1_found = any(
            t.content == "Main Heading" and hasattr(t, 'level') and t.level == 1
            for t in result.texts
        )
        assert heading1_found, "Should identify heading level 1"

    @pytest.mark.asyncio
    async def test_parser_docx_heading_level_2_recognition(self, sample_docx_with_content: Path) -> None:
        """Test that heading level 2 is correctly recognized."""
        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(sample_docx_with_content)

        # Find heading with level 2
        heading2_found = any(
            t.content == "Sub Heading" and hasattr(t, 'level') and t.level == 2
            for t in result.texts
        )
        assert heading2_found, "Should identify heading level 2"

    @pytest.mark.asyncio
    async def test_parser_docx_heading_level_3_recognition(self, sample_docx_with_content: Path) -> None:
        """Test that heading level 3 is correctly recognized."""
        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(sample_docx_with_content)

        # Find heading with level 3
        heading3_found = any(
            t.content == "Sub-sub Heading" and hasattr(t, 'level') and t.level == 3
            for t in result.texts
        )
        assert heading3_found, "Should identify heading level 3"

    @pytest.mark.asyncio
    async def test_parser_docx_paragraph_extraction(self, sample_docx_with_content: Path) -> None:
        """Test that paragraph content is correctly extracted."""
        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(sample_docx_with_content)

        # Check for expected paragraph content
        text_content = " ".join([t.content for t in result.texts])
        assert "sample paragraph" in text_content.lower(), "Should extract first paragraph"
        assert "Second paragraph" in text_content, "Should extract second paragraph"

    @pytest.mark.asyncio
    async def test_parser_docx_table_extraction(self, sample_docx_with_content: Path) -> None:
        """Test that table structure is correctly extracted."""
        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(sample_docx_with_content)

        assert len(result.tables) == 1, "Should extract exactly one table"

        table = result.tables[0]
        assert len(table.headers) == 2, "Table should have 2 headers"
        # TableElement uses data field (list of lists), not rows
        assert len(table.data) == 3, "Table should have 3 rows including header"

        # Verify header content
        assert "Column 1" in table.headers[0], "Should extract first header"
        assert "Column 2" in table.headers[1], "Should extract second header"

        # Verify row content from data field
        row_content = " ".join([cell for row in table.data for cell in row])
        assert "Data 1" in row_content, "Should extract first row data"
        assert "Data 4" in row_content, "Should extract last row data"

    @pytest.mark.asyncio
    async def test_parser_docx_empty_document(self, empty_docx: Path) -> None:
        """Test parsing empty DOCX."""
        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(empty_docx)

        assert isinstance(result, UnifiedDocument)
        assert result.fingerprint != "", "Should still generate fingerprint for empty file"
        assert result.parse_method == "python-docx"
        # Empty document should have no texts or tables
        assert len(result.texts) == 0, "Empty DOCX should have no texts"
        assert len(result.tables) == 0, "Empty DOCX should have no tables"

    @pytest.mark.asyncio
    async def test_parser_docx_nonexistent_file(self, tmp_path: Path) -> None:
        """Test parsing non-existent DOCX file raises appropriate error."""
        parser = DocxParser()
        nonexistent_file = tmp_path / "nonexistent.docx"

        with pytest.raises(FileNotFoundError):
            await parser.parse(nonexistent_file)

    @pytest.mark.asyncio
    async def test_parser_docx_multiple_headings_same_level(self, tmp_path: Path) -> None:
        """Test parsing DOCX with multiple headings of same level."""
        doc = Document()

        # Add multiple level 1 headings
        doc.add_heading('First Heading', level=1)
        doc.add_paragraph('Content for first heading.')
        doc.add_heading('Second Heading', level=1)
        doc.add_paragraph('Content for second heading.')

        file_path = tmp_path / "multiple_headings.docx"
        doc.save(file_path)

        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(file_path)

        # Should extract both headings
        heading1_count = sum(
            1 for t in result.texts
            if t.content == "First Heading" and hasattr(t, 'level') and t.level == 1
        )
        heading2_count = sum(
            1 for t in result.texts
            if t.content == "Second Heading" and hasattr(t, 'level') and t.level == 1
        )

        assert heading1_count == 1, "Should extract first heading"
        assert heading2_count == 1, "Should extract second heading"

    @pytest.mark.asyncio
    async def test_parser_docx_multiple_tables(self, tmp_path: Path) -> None:
        """Test parsing DOCX with multiple tables."""
        doc = Document()

        # First table
        table1 = doc.add_table(rows=2, cols=2)
        table1.rows[0].cells[0].text = "Header 1"
        table1.rows[0].cells[1].text = "Header 2"
        table1.rows[1].cells[0].text = "Data 1"
        table1.rows[1].cells[1].text = "Data 2"

        # Paragraph between tables
        doc.add_paragraph("Text between tables.")

        # Second table
        table2 = doc.add_table(rows=2, cols=2)
        table2.rows[0].cells[0].text = "Header A"
        table2.rows[0].cells[1].text = "Header B"
        table2.rows[1].cells[0].text = "Data A"
        table2.rows[1].cells[1].text = "Data B"

        file_path = tmp_path / "multiple_tables.docx"
        doc.save(file_path)

        parser = DocxParser()
        result: UnifiedDocument = await parser.parse(file_path)

        assert len(result.tables) == 2, "Should extract both tables"

    @pytest.mark.asyncio
    async def test_parser_docx_fingerprint_uniqueness(self, sample_docx_with_content: Path) -> None:
        """Test that fingerprint is unique and consistent for same file."""
        parser = DocxParser()

        # Parse same file twice
        result1: UnifiedDocument = await parser.parse(sample_docx_with_content)
        result2: UnifiedDocument = await parser.parse(sample_docx_with_content)

        # Fingerprints should be identical for same file
        assert result1.fingerprint == result2.fingerprint, "Fingerprint should be consistent"

        # Fingerprint should not be empty
        assert len(result1.fingerprint) > 0, "Fingerprint should not be empty"